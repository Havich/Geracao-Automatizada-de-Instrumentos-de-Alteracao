import streamlit as st
import streamlit.components.v1 as components
import sqlite3
import hashlib
import io
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==========================================
# CONFIGURAÇÃO DE LAYOUT E CSS
# ==========================================
st.set_page_config(page_title="GAIA", page_icon= "logo.png", layout="centered")

st.markdown("""
    <style>
    header[data-testid="stHeader"] { display: none !important; }
    .block-container { padding-top: 5rem; padding-bottom: 1rem; }
    .stApp { background-color: #FFFFFF; }
    h1, h2, h3 { color: #006400 !important; text-align: center; font-family: 'Arial'; margin-top: 0px; margin-bottom: 5px;}
    [data-testid="stImage"] { display: flex; justify-content: center; }
    button[kind="primary"], button[data-testid="baseButton-primary"],
    button[kind="primaryFormSubmit"], button[data-testid="baseButton-primaryFormSubmit"] {
        background-color: #28a745 !important; border-color: #28a745 !important; color: white !important;
    }
    button[kind="secondary"], button[data-testid="baseButton-secondary"],
    button[kind="secondaryFormSubmit"], button[data-testid="baseButton-secondaryFormSubmit"] {
        background-color: #6c757d !important; border-color: #6c757d !important; color: white !important;
    }
    .stButton>button, .stFormSubmitButton>button {
        border-radius: 8px; width: 100%; height: 38px; font-weight: bold; transition: 0.3s;
    }
    hr { margin: 8px 0px; border-color: #e0e0e0; }
    .title-separator { border-top: 2px solid #006400; margin-top: 20px; margin-bottom: 10px;}
    </style>
""", unsafe_allow_html=True)

def cabecalho():
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        try:
            st.image("logo.png", width=400)
        except:
            st.markdown("<h2>🌿 GAIA</h2>", unsafe_allow_html=True)

# ==========================================
# BANCO DE DADOS
# ==========================================
def get_db():
    conn = sqlite3.connect('gaia.db')
    conn.execute("PRAGMA foreign_keys = ON;")
    return conn

def inicializar_banco():
    with get_db() as conn:
        conn.execute("CREATE TABLE IF NOT EXISTS tb_perfil_acesso (id INTEGER PRIMARY KEY, nome_perfil TEXT UNIQUE NOT NULL)")
        conn.execute("""CREATE TABLE IF NOT EXISTS tb_usuario (
            id INTEGER PRIMARY KEY AUTOINCREMENT, nome TEXT, login TEXT UNIQUE,
            senha_hash TEXT, id_perfil INTEGER,
            FOREIGN KEY (id_perfil) REFERENCES tb_perfil_acesso (id))""")
        conn.execute("""CREATE TABLE IF NOT EXISTS tb_empresa (
            id INTEGER PRIMARY KEY AUTOINCREMENT, cnpj TEXT UNIQUE, nome_empresarial TEXT,
            nome_fantasia TEXT,
            nire TEXT, cidade TEXT, uf TEXT, cep TEXT, logradouro TEXT, numero TEXT,
            bairro TEXT, complemento TEXT)""")
        conn.execute("""CREATE TABLE IF NOT EXISTS tb_socio (
            id INTEGER PRIMARY KEY AUTOINCREMENT, id_empresa INTEGER, nome TEXT,
            nacionalidade TEXT, cpf TEXT, rg TEXT, orgao_emissor TEXT,
            estado_civil TEXT, data_nasc TEXT, profissao TEXT, cep TEXT,
            logradouro TEXT, numero TEXT, complemento TEXT, bairro TEXT, cidade TEXT, uf TEXT,
            FOREIGN KEY (id_empresa) REFERENCES tb_empresa (id) ON DELETE CASCADE)""")
        conn.execute("""CREATE TABLE IF NOT EXISTS tb_documento (
            id INTEGER PRIMARY KEY AUTOINCREMENT, id_empresa INTEGER, id_usuario INTEGER,
            data_geracao DATETIME DEFAULT CURRENT_TIMESTAMP, tipo_documento TEXT, status TEXT DEFAULT 'Gerado',
            FOREIGN KEY (id_empresa) REFERENCES tb_empresa (id) ON DELETE CASCADE,
            FOREIGN KEY (id_usuario) REFERENCES tb_usuario (id))""")
        conn.commit()

        # Migração: adiciona colunas que podem não existir em bancos antigos
        colunas_empresa = [r[1] for r in conn.execute("PRAGMA table_info(tb_empresa)").fetchall()]
        if 'nome_fantasia' not in colunas_empresa:
            conn.execute("ALTER TABLE tb_empresa ADD COLUMN nome_fantasia TEXT")
            conn.commit()

        if conn.execute("SELECT COUNT(*) FROM tb_perfil_acesso").fetchone()[0] == 0:
            conn.execute("INSERT INTO tb_perfil_acesso (nome_perfil) VALUES ('Admin'), ('Operador')")
            conn.commit()

        if conn.execute("SELECT COUNT(*) FROM tb_usuario").fetchone()[0] == 0:
            h = hashlib.sha256("1234".encode()).hexdigest()
            conn.execute("INSERT INTO tb_usuario (nome, login, senha_hash, id_perfil) VALUES (?,?,?,1)", ("Administrador", "admin", h))
            conn.commit()

inicializar_banco()

# ==========================================
# ESTADO E NAVEGAÇÃO
# ==========================================
defaults = {
    'pagina': 'login',
    'empresa_id': None,
    'socio_id': None,
    'user_id': None,
    'evento_tipo': None,
    'evento_dados': {}
}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

def gerar_hash(s):
    return hashlib.sha256(s.encode()).hexdigest()

def nav(p):
    st.session_state.pagina = p
    st.rerun()

# ==========================================
# FORMULÁRIO REUTILIZÁVEL DE SÓCIO
# ==========================================
def _form_socio(titulo, defaults=None):
    d = defaults or {}
    estados_civis = ["", "Solteiro(a)", "Casado(a)", "Divorciado(a)", "Viúvo(a)", "União Estável"]
    cabecalho()
    st.subheader(titulo)

    with st.form("f_socio"):
        nome = st.text_input("NOME COMPLETO*", value=d.get('nome') or '')
        c1, c2 = st.columns(2)
        nac = c1.text_input("NACIONALIDADE*", value=d.get('nacionalidade') or 'BRASILEIRO(A)')
        profissao = c2.text_input("PROFISSÃO", value=d.get('profissao') or '')
        c3, c4 = st.columns(2)
        cpf = c3.text_input("CPF*", value=d.get('cpf') or '')
        # Máscara JS removida para evitar o bug de apagar a data
        data_nasc = c4.text_input("DATA NASCIMENTO (DD/MM/AAAA)", value=d.get('data_nasc') or '')
        c5, c6 = st.columns(2)
        rg = c5.text_input("RG", value=d.get('rg') or '')
        orgao = c6.text_input("ÓRGÃO EMISSOR", value=d.get('orgao_emissor') or '')
        ec_valor = d.get('estado_civil') or ""
        ec_idx = estados_civis.index(ec_valor) if ec_valor in estados_civis else 0
        estado_civil = st.selectbox("ESTADO CIVIL", estados_civis, index=ec_idx)
        st.markdown("**Endereço**")
        c7, c8 = st.columns([3, 1])
        logradouro = c7.text_input("LOGRADOURO*", value=d.get('logradouro') or '')
        numero = c8.text_input("NÚMERO*", value=d.get('numero') or '')
        c9, c10 = st.columns(2)
        complemento = c9.text_input("COMPLEMENTO", value=d.get('complemento') or '')
        bairro = c10.text_input("BAIRRO*", value=d.get('bairro') or '')
        c11, c12, c13 = st.columns([3, 1, 2])
        cidade = c11.text_input("CIDADE*", value=d.get('cidade') or '')
        uf = c12.text_input("UF*", value=d.get('uf') or '')
        cep = c13.text_input("CEP*", value=d.get('cep') or '')
        b1, b2 = st.columns(2)
        btn_v = b1.form_submit_button("⬅️ CANCELAR", type="secondary", use_container_width=True)
        btn_ok = b2.form_submit_button("💾 CONFIRMAR", type="primary", use_container_width=True)

    return btn_v, btn_ok, dict(
        nome=(nome or "").upper(), nacionalidade=(nac or "").upper(), cpf=cpf, rg=rg,
        orgao_emissor=(orgao or "").upper(), estado_civil=estado_civil, data_nasc=data_nasc,
        profissao=(profissao or "").upper(), logradouro=(logradouro or "").upper(),
        numero=numero, complemento=(complemento or "").upper(), bairro=(bairro or "").upper(),
        cidade=(cidade or "").upper(), uf=(uf or "").upper(), cep=cep
    )

# ==========================================
# GERAÇÃO DO DOCUMENTO WORD
# ==========================================
def gerar_doc_alteracao(empresa, socios, evento_tipo, evento_dados):
    doc = Document()

    # Margens
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    def add_paragrafo(texto, bold=False, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(texto)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = 'Arial'
        return p

    # Título
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo.paragraph_format.space_after = Pt(12)
    run_titulo = titulo.add_run(f"ALTERAÇÃO CONTRATUAL {empresa['nome_empresarial']}")
    run_titulo.bold = True
    run_titulo.font.size = Pt(14)
    run_titulo.font.name = 'Arial'

    # Qualificação dos sócios
    for i, s in enumerate(socios):
        partes = []
        partes.append(s.get('nome', ''))
        if s.get('nacionalidade'):
            partes.append(f"nacionalidade {s['nacionalidade']}")
        if s.get('estado_civil'):
            partes.append(s['estado_civil'])
        if s.get('data_nasc'):
            partes.append(f"nascido em {s['data_nasc']}")
        if s.get('profissao'):
            partes.append(f"profissão: {s['profissao']}")
        if s.get('cpf'):
            partes.append(f"nº do CPF: {s['cpf']}")
        if s.get('rg'):
            partes.append(f"identidade: {s['rg']}")
        if s.get('orgao_emissor'):
            partes.append(f"órgão expedidor: {s['orgao_emissor']}")

        end_partes = []
        if s.get('logradouro'):
            end_partes.append(s['logradouro'])
        if s.get('complemento'):
            end_partes.append(s['complemento'])
        if s.get('numero'):
            end_partes.append(f"número {s['numero']}")
        if s.get('bairro'):
            end_partes.append(f"bairro {s['bairro']}")
        if s.get('cidade') and s.get('uf'):
            end_partes.append(f"município {s['cidade']} - {s['uf']}")
        if s.get('cep'):
            end_partes.append(f"CEP: {s['cep']}")

        if end_partes:
            partes.append(f"RESIDENTE E DOMICILIADO no(a): {', '.join(end_partes)}")

        texto_socio = ', '.join(partes) + ','
        add_paragrafo(texto_socio)

    # Qualificação da empresa
    emp_partes = [
        f"Sócio(s) da sociedade limitada {empresa['nome_empresarial']}",
    ]
    end_emp = []
    if empresa.get('logradouro'):
        end_emp.append(empresa['logradouro'])
    if empresa.get('numero'):
        end_emp.append(f"número {empresa['numero']}")
    if empresa.get('bairro'):
        end_emp.append(f"bairro {empresa['bairro']}")
    if empresa.get('complemento'):
        end_emp.append(empresa['complemento'])
    if empresa.get('cidade') and empresa.get('uf'):
        end_emp.append(f"município {empresa['cidade']} - {empresa['uf']}")
    if empresa.get('cep'):
        end_emp.append(f"CEP: {empresa['cep']}")
    if end_emp:
        emp_partes.append(f"sediada na {', '.join(end_emp)}")
    if empresa.get('nire'):
        emp_partes.append(f"com seu contrato social arquivado nessa Junta Comercial")
    if empresa.get('cnpj'):
        emp_partes.append(f"devidamente inscrita no CNPJ sob o nº {empresa['cnpj']}")
    emp_partes.append("resolvem:")

    add_paragrafo(', '.join(emp_partes))

    # Cláusula conforme evento
    clausula_num = 1

    if evento_tipo == 'mudanca_endereco':
        novo_end = evento_dados
        partes_end = []
        if novo_end.get('logradouro'):
            partes_end.append(novo_end['logradouro'])
        if novo_end.get('numero'):
            partes_end.append(f"número {novo_end['numero']}")
        if novo_end.get('bairro'):
            partes_end.append(f"bairro {novo_end['bairro']}")
        if novo_end.get('complemento'):
            partes_end.append(novo_end['complemento'])
        if novo_end.get('cidade') and novo_end.get('uf'):
            partes_end.append(f"município {novo_end['cidade']} - {novo_end['uf']}")
        if novo_end.get('cep'):
            partes_end.append(f"CEP: {novo_end['cep']}")

        texto_clausula = (
            f"Cláusula {_ordinal(clausula_num)} – A sociedade passará a ter sua sede no(a): "
            f"{', '.join(partes_end)}."
        )
        add_paragrafo(texto_clausula, bold=True if clausula_num == 1 else False)
        clausula_num += 1

    elif evento_tipo == 'mudanca_nome':
        novo_nome = evento_dados.get('nome_empresarial', '')
        texto_clausula = (
            f"Cláusula {_ordinal(clausula_num)} – A sociedade passará a denominar-se: {novo_nome}."
        )
        add_paragrafo(texto_clausula)
        clausula_num += 1

    elif evento_tipo == 'mudanca_nome_fantasia':
        novo_nf = evento_dados.get('nome_fantasia', '')
        texto_clausula = (
            f"Cláusula {_ordinal(clausula_num)} – A sociedade passará a ter o TÍTULO DO ESTABELECIMENTO "
            f"(NOME FANTASIA): {novo_nf}."
        )
        add_paragrafo(texto_clausula)
        clausula_num += 1

    # Cláusula final padrão
    add_paragrafo(
        f"Cláusula {_ordinal(clausula_num)} – As demais cláusulas e condições do contrato social "
        "permanecem inalteradas."
    )

    add_paragrafo("")

    # Local/data
    add_paragrafo(
        f"{empresa.get('cidade', '_____________')}, _____ de ________________ de _______.",
        align=WD_ALIGN_PARAGRAPH.CENTER
    )

    add_paragrafo("")

    # Assinaturas
    for s in socios:
        add_paragrafo("_" * 50, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
        add_paragrafo(s.get('nome', ''), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def _ordinal(n):
    nomes = {1: "Primeira", 2: "Segunda", 3: "Terceira", 4: "Quarta", 5: "Quinta"}
    return nomes.get(n, str(n) + "ª")

# ==========================================
# TELAS
# ==========================================

def tela_login():
    cabecalho()
    st.subheader("Acesso ao sistema")
    with st.form("form_login"):
        login = st.text_input("Login")
        senha = st.text_input("Senha", type="password")
        btn_e = st.form_submit_button("ENTRAR", type="primary", use_container_width=True)
        btn_cad = st.form_submit_button("Ainda não tem conta? CADASTRE-SE", type="secondary", use_container_width=True)
    if btn_e:
        with get_db() as conn:
            user = conn.execute(
                "SELECT id FROM tb_usuario WHERE login=? AND senha_hash=?",
                (login, gerar_hash(senha))
            ).fetchone()
        if user:
            st.session_state.user_id = user[0]
            nav('menu')
        else:
            st.error("Login ou senha incorretos.")
    if btn_cad:
        nav('registrar')

def tela_registrar():
    cabecalho()
    st.subheader("Criar nova conta")
    with st.form("form_reg"):
        u = st.text_input("Usuário*")
        s1 = st.text_input("Senha*", type="password")
        s2 = st.text_input("Confirmar senha*", type="password")
        c1, c2 = st.columns(2)
        btn_voltar = c1.form_submit_button("⬅️ CANCELAR", type="secondary")
        btn_confirmar = c2.form_submit_button("CONFIRMAR", type="primary")

    if btn_voltar:
        nav('login')

    if btn_confirmar:
        if not u:
            st.error("Informe um nome de usuário.")
        elif s1 != s2:
            st.error("As senhas não conferem.")
        elif len(s1) < 4:
            st.error("A senha deve ter pelo menos 4 caracteres.")
        else:
            try:
                with get_db() as conn:
                    conn.execute(
                        "INSERT INTO tb_usuario (nome, login, senha_hash, id_perfil) VALUES (?,?,?,2)",
                        (u, u, gerar_hash(s1))
                    )
                    conn.commit()
                st.success(f"Usuário **{u}** criado com sucesso! Faça login para continuar.")
                import time; time.sleep(1)
                nav('login')
            except sqlite3.IntegrityError:
                st.error(f"O usuário '{u}' já existe. Escolha outro nome.")

def tela_menu():
    cabecalho()
    st.subheader("Bem vindo ao GAIA!")
    if st.button("➕ Nova empresa / Alteração", type="primary", use_container_width=True):
        nav('cad_empresa')
    if st.button("🔍 Buscar empresa cadastrada", type="primary", use_container_width=True):
        nav('buscar_empresa')
    if st.button("📄 Histórico geral de documentos", type="secondary", use_container_width=True):
        nav('historico')
    st.divider()
    if st.button("⬅️ Sair da conta", type="secondary", use_container_width=True):
        st.session_state.user_id = None
        nav('login')

def tela_cad_empresa():
    cabecalho()
    st.subheader("DADOS DA EMPRESA")
    with st.form("f_emp"):
        cnpj = st.text_input("CNPJ*")
        nome = st.text_input("NOME EMPRESARIAL*")
        nire = st.text_input("NIRE (opcional)")
        st.markdown("**Endereço**")
        c7, c8 = st.columns([3, 1])
        logradouro = c7.text_input("LOGRADOURO*")
        numero = c8.text_input("NÚMERO*")
        c9, c10 = st.columns(2)
        complemento = c9.text_input("COMPLEMENTO")
        bairro = c10.text_input("BAIRRO*")
        c11, c12, c13 = st.columns([3, 1, 2])
        cidade = c11.text_input("CIDADE*")
        uf = c12.text_input("UF*")
        cep = c13.text_input("CEP*")
        c1, c2 = st.columns(2)
        btn_voltar = c1.form_submit_button("⬅️ CANCELAR", type="secondary")
        btn_seguir = c2.form_submit_button("SEGUIR ➡️", type="primary")

    if btn_voltar:
        nav('menu')

    if btn_seguir:
        erros = []
        if not cnpj: erros.append("CNPJ")
        if not nome: erros.append("Nome Empresarial")
        if not logradouro: erros.append("Logradouro")
        if not numero: erros.append("Número")
        if not bairro: erros.append("Bairro")
        if not cidade: erros.append("Cidade")
        if not uf: erros.append("UF")
        if not cep: erros.append("CEP")
        if erros:
            st.error(f"Preencha os campos obrigatórios: {', '.join(erros)}.")
        else:
            try:
                with get_db() as conn:
                    cur = conn.cursor()
                    cur.execute(
                        "INSERT INTO tb_empresa (cnpj, nome_empresarial, nire, logradouro, numero, complemento, bairro, cidade, uf, cep) VALUES (?,?,?,?,?,?,?,?,?,?)",
                        (cnpj, nome.upper(), nire, logradouro.upper(), numero, complemento.upper(), bairro.upper(), cidade.upper(), uf.upper(), cep)
                    )
                    st.session_state.empresa_id = cur.lastrowid
                    conn.commit()
                nav('lista_socios')
            except sqlite3.IntegrityError:
                st.error(f"Já existe uma empresa com o CNPJ '{cnpj}'. Use 'Buscar empresa' para acessá-la.")

def tela_buscar_empresa():
    cabecalho()
    st.subheader("SELECIONE UMA EMPRESA")
    with get_db() as conn:
        empresas = conn.execute("SELECT id, cnpj, nome_empresarial FROM tb_empresa").fetchall()
    if not empresas:
        st.info("Nenhuma empresa cadastrada.")
        if st.button("⬅️ VOLTAR AO MENU", type="secondary"):
            nav('menu')
    else:
        opcoes = {f"{e[1]} - {e[2]}": e[0] for e in empresas}
        escolha = st.selectbox("Empresas:", list(opcoes.keys()))
        if st.button("ACESSAR EMPRESA", type="primary", use_container_width=True):
            st.session_state.empresa_id = opcoes[escolha]
            nav('lista_socios')
        if st.button("⬅️ VOLTAR AO MENU", type="secondary", use_container_width=True):
            nav('menu')

def tela_editar_empresa():
    cabecalho()
    st.subheader("EDITAR EMPRESA")
    with get_db() as conn:
        e = conn.execute(
            "SELECT cnpj, nome_empresarial, nire, logradouro, numero, complemento, bairro, cidade, uf, cep FROM tb_empresa WHERE id=?",
            (st.session_state.empresa_id,)
        ).fetchone()

    with st.form("f_edit_emp"):
        cnpj = st.text_input("CNPJ*", value=e[0] or '')
        nome = st.text_input("NOME EMPRESARIAL*", value=e[1] or '')
        nire = st.text_input("NIRE", value=e[2] or '')
        st.markdown("**Endereço**")
        c7, c8 = st.columns([3, 1])
        logradouro = c7.text_input("LOGRADOURO*", value=e[3] or '')
        numero = c8.text_input("NÚMERO*", value=e[4] or '')
        c9, c10 = st.columns(2)
        complemento = c9.text_input("COMPLEMENTO", value=e[5] or '')
        bairro = c10.text_input("BAIRRO*", value=e[6] or '')
        c11, c12, c13 = st.columns([3, 1, 2])
        cidade = c11.text_input("CIDADE*", value=e[7] or '')
        uf = c12.text_input("UF*", value=e[8] or '')
        cep = c13.text_input("CEP*", value=e[9] or '')
        c1, c2 = st.columns(2)
        btn_v = c1.form_submit_button("⬅️ CANCELAR", type="secondary")
        btn_ok = c2.form_submit_button("💾 SALVAR", type="primary")

    if btn_v:
        nav('lista_socios')
    if btn_ok:
        erros = []
        if not cnpj: erros.append("CNPJ")
        if not nome: erros.append("Nome Empresarial")
        if not logradouro: erros.append("Logradouro")
        if not numero: erros.append("Número")
        if not bairro: erros.append("Bairro")
        if not cidade: erros.append("Cidade")
        if not uf: erros.append("UF")
        if not cep: erros.append("CEP")
        if erros:
            st.error(f"Preencha os campos obrigatórios: {', '.join(erros)}.")
        else:
            with get_db() as conn:
                conn.execute(
                    "UPDATE tb_empresa SET cnpj=?, nome_empresarial=?, nire=?, logradouro=?, numero=?, complemento=?, bairro=?, cidade=?, uf=?, cep=? WHERE id=?",
                    (cnpj, nome.upper(), nire, logradouro.upper(), numero, complemento.upper(), bairro.upper(), cidade.upper(), uf.upper(), cep, st.session_state.empresa_id)
                )
                conn.commit()
            st.success("Empresa atualizada com sucesso!")
            nav('lista_socios')

def tela_socios():
    cabecalho()
    with get_db() as conn:
        emp = conn.execute(
            "SELECT nome_empresarial FROM tb_empresa WHERE id=?",
            (st.session_state.empresa_id,)
        ).fetchone()
        socios = conn.execute(
            "SELECT id, nome, cpf FROM tb_socio WHERE id_empresa=?",
            (st.session_state.empresa_id,)
        ).fetchall()

    # --- Confirmação exclusão empresa ---
    if st.session_state.get('_confirm_excluir_empresa'):
        st.warning(f"⚠️ Tem certeza que deseja excluir a empresa **{emp[0]}** e todos os seus sócios? Esta ação não pode ser desfeita.")
        ca, cb = st.columns(2)
        if ca.button("❌ Cancelar", type="secondary", use_container_width=True):
            st.session_state.pop('_confirm_excluir_empresa', None)
            st.rerun()
        if cb.button("🗑️ Sim, excluir empresa", type="primary", use_container_width=True):
            with get_db() as conn:
                conn.execute("DELETE FROM tb_empresa WHERE id=?", (st.session_state.empresa_id,))
                conn.commit()
            st.session_state.empresa_id = None
            st.session_state.pop('_confirm_excluir_empresa', None)
            nav('menu')
        return

    # --- Confirmação exclusão sócio ---
    if st.session_state.get('_confirm_excluir_socio'):
        sid, snome = st.session_state['_confirm_excluir_socio']
        st.warning(f"⚠️ Tem certeza que deseja excluir o sócio **{snome}**?")
        ca, cb = st.columns(2)
        if ca.button("❌ Cancelar", type="secondary", use_container_width=True):
            st.session_state.pop('_confirm_excluir_socio', None)
            st.rerun()
        if cb.button("🗑️ Sim, excluir sócio", type="primary", use_container_width=True):
            with get_db() as conn:
                conn.execute("DELETE FROM tb_socio WHERE id=?", (sid,))
                conn.commit()
            st.session_state.pop('_confirm_excluir_socio', None)
            st.rerun()
        return

    col_n, col_ed, col_ex = st.columns([0.6, 0.2, 0.2], vertical_alignment="center")
    col_n.markdown(f"### {emp[0]}")
    if col_ed.button("✏️ Editar", key="ed_emp"):
        nav('editar_empresa')
    if col_ex.button("🗑️ Excluir", key="ex_emp"):
        st.session_state['_confirm_excluir_empresa'] = True
        st.rerun()

    st.markdown("<div class='title-separator'></div>", unsafe_allow_html=True)
    st.markdown("<h4 style='color:#006400;text-align:center;'>Quadro Societário</h4>", unsafe_allow_html=True)

    for s in socios:
        c1, c2, c3 = st.columns([0.6, 0.2, 0.2], vertical_alignment="center")
        c1.markdown(f"☑ **{s[1]}** *(CPF: {s[2]})*")
        if c2.button("✏️ Editar", key=f"ed_s_{s[0]}", type="secondary"):
            st.session_state.socio_id = s[0]
            nav('editar_socio')
        if c3.button("🗑️ Excluir", key=f"ex_s_{s[0]}", type="secondary"):
            st.session_state['_confirm_excluir_socio'] = (s[0], s[1])
            st.rerun()

    st.markdown("<br>", unsafe_allow_html=True)
    b1, b2, b3 = st.columns([1, 2, 1])
    with b2:
        if st.button("➕ Novo Sócio", type="primary", use_container_width=True):
            nav('cad_socio')

    st.markdown("<br>", unsafe_allow_html=True)
    if st.button("📄 Histórico de Documentos desta Empresa", type="secondary", use_container_width=True):
        nav('historico_empresa')

    st.markdown("<hr>", unsafe_allow_html=True)
    n1, n2 = st.columns(2)
    if n1.button("⬅️ VOLTAR AO MENU", type="secondary"):
        st.session_state.empresa_id = None
        nav('menu')
    if n2.button("AVANÇAR PARA EVENTOS ➡️", type="primary"):
        nav('escolher_evento')

def tela_cad_socio():
    btn_v, btn_ok, dados = _form_socio("CADASTRAR NOVO SÓCIO")
    if btn_v:
        nav('lista_socios')
    if btn_ok:
        erros = []
        if not dados['nome']: erros.append("Nome completo")
        if not dados['cpf']: erros.append("CPF")
        if not dados['logradouro']: erros.append("Logradouro")
        if not dados['numero']: erros.append("Número")
        if not dados['bairro']: erros.append("Bairro")
        if not dados['cidade']: erros.append("Cidade")
        if not dados['uf']: erros.append("UF")
        if not dados['cep']: erros.append("CEP")
        if erros:
            st.error(f"Preencha os campos obrigatórios: {', '.join(erros)}.")
        else:
            with get_db() as conn:
                conn.execute(
                    """INSERT INTO tb_socio
                    (id_empresa, nome, nacionalidade, cpf, rg, orgao_emissor, estado_civil,
                    data_nasc, profissao, logradouro, numero, complemento, bairro, cidade, uf, cep)
                    VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                    (
                        st.session_state.empresa_id,
                        dados['nome'], dados['nacionalidade'], dados['cpf'],
                        dados['rg'], dados['orgao_emissor'], dados['estado_civil'],
                        dados['data_nasc'], dados['profissao'], dados['logradouro'],
                        dados['numero'], dados['complemento'], dados['bairro'],
                        dados['cidade'], dados['uf'], dados['cep']
                    )
                )
                conn.commit()
            st.success(f"Sócio **{dados['nome']}** cadastrado com sucesso!")
            st.session_state['_redir_socios'] = True

    if st.session_state.pop('_redir_socios', False):
        nav('lista_socios')

def tela_editar_socio():
    with get_db() as conn:
        s = conn.execute(
            """SELECT nome, nacionalidade, cpf, rg, orgao_emissor, estado_civil,
            data_nasc, profissao, logradouro, numero, complemento, bairro, cidade, uf, cep
            FROM tb_socio WHERE id=?""",
            (st.session_state.socio_id,)
        ).fetchone()

    defaults = {
        'nome': s[0], 'nacionalidade': s[1], 'cpf': s[2], 'rg': s[3],
        'orgao_emissor': s[4], 'estado_civil': s[5], 'data_nasc': s[6],
        'profissao': s[7], 'logradouro': s[8], 'numero': s[9],
        'complemento': s[10], 'bairro': s[11], 'cidade': s[12], 'uf': s[13], 'cep': s[14]
    }
    btn_v, btn_ok, dados = _form_socio("EDITAR SÓCIO", defaults)
    if btn_v:
        nav('lista_socios')
    if btn_ok:
        with get_db() as conn:
            conn.execute(
                """UPDATE tb_socio SET nome=?, nacionalidade=?, cpf=?, rg=?, orgao_emissor=?,
                estado_civil=?, data_nasc=?, profissao=?, logradouro=?, numero=?,
                complemento=?, bairro=?, cidade=?, uf=?, cep=? WHERE id=?""",
                (
                    dados['nome'], dados['nacionalidade'], dados['cpf'], dados['rg'],
                    dados['orgao_emissor'], dados['estado_civil'], dados['data_nasc'],
                    dados['profissao'], dados['logradouro'], dados['numero'],
                    dados['complemento'], dados['bairro'], dados['cidade'],
                    dados['uf'], dados['cep'], st.session_state.socio_id
                )
            )
            conn.commit()
        nav('lista_socios')

# ==========================================
# EVENTOS
# ==========================================

def tela_escolher_evento():
    cabecalho()
    st.subheader("ESCOLHA O TIPO DE ALTERAÇÃO")
    st.markdown("<br>", unsafe_allow_html=True)

    if st.button("🏠 Mudança de Endereço da Sede", type="primary", use_container_width=True):
        st.session_state.evento_tipo = 'mudanca_endereco'
        nav('evento_endereco')

    if st.button("📝 Mudança de Nome Empresarial", type="primary", use_container_width=True):
        st.session_state.evento_tipo = 'mudanca_nome'
        nav('evento_nome')

    if st.button("🏷️ Mudança de Nome Fantasia", type="primary", use_container_width=True):
        st.session_state.evento_tipo = 'mudanca_nome_fantasia'
        nav('evento_nome_fantasia')

    st.markdown("<br>", unsafe_allow_html=True)
    if st.button("⬅️ VOLTAR", type="secondary", use_container_width=True):
        nav('lista_socios')

def tela_evento_endereco():
    cabecalho()
    st.subheader("MUDANÇA DE ENDEREÇO DA SEDE")
    with st.form("f_end"):
        c7, c8 = st.columns([3, 1])
        logradouro = c7.text_input("LOGRADOURO*")
        numero = c8.text_input("NÚMERO*")
        c9, c10 = st.columns(2)
        complemento = c9.text_input("COMPLEMENTO")
        bairro = c10.text_input("BAIRRO*")
        c11, c12, c13 = st.columns([3, 1, 2])
        cidade = c11.text_input("CIDADE*")
        uf = c12.text_input("UF*")
        cep = c13.text_input("CEP*")
        c1, c2 = st.columns(2)
        btn_v = c1.form_submit_button("⬅️ CANCELAR", type="secondary")
        btn_ok = c2.form_submit_button("📄 GERAR DOCUMENTO", type="primary")

    if btn_v:
        nav('escolher_evento')
    if btn_ok:
        if not logradouro or not cidade or not uf:
            st.error("Preencha os campos obrigatórios.")
        else:
            st.session_state.evento_dados = dict(
                logradouro=logradouro.upper(), numero=numero,
                complemento=complemento.upper(), bairro=bairro.upper(),
                cidade=cidade.upper(), uf=uf.upper(), cep=cep
            )
            nav('gerar_documento')

def tela_evento_nome():
    cabecalho()
    st.subheader("MUDANÇA DE NOME EMPRESARIAL")
    with st.form("f_nome"):
        novo_nome = st.text_input("NOVO NOME EMPRESARIAL*")
        c1, c2 = st.columns(2)
        btn_v = c1.form_submit_button("⬅️ CANCELAR", type="secondary")
        btn_ok = c2.form_submit_button("📄 GERAR DOCUMENTO", type="primary")
    if btn_v:
        nav('escolher_evento')
    if btn_ok:
        if not novo_nome:
            st.error("Informe o novo nome.")
        else:
            st.session_state.evento_dados = {'nome_empresarial': novo_nome.upper()}
            nav('gerar_documento')

def tela_evento_nome_fantasia():
    cabecalho()
    st.subheader("MUDANÇA DE NOME FANTASIA")
    with st.form("f_nf"):
        novo_nf = st.text_input("NOVO NOME FANTASIA*")
        c1, c2 = st.columns(2)
        btn_v = c1.form_submit_button("⬅️ CANCELAR", type="secondary")
        btn_ok = c2.form_submit_button("📄 GERAR DOCUMENTO", type="primary")
    if btn_v:
        nav('escolher_evento')
    if btn_ok:
        if not novo_nf:
            st.error("Informe o nome fantasia.")
        else:
            st.session_state.evento_dados = {'nome_fantasia': novo_nf.upper()}
            nav('gerar_documento')

def tela_gerar_documento():
    cabecalho()
    st.subheader("📄 DOCUMENTO GERADO")

    with get_db() as conn:
        emp_row = conn.execute(
            "SELECT cnpj, nome_empresarial, nome_fantasia, nire, logradouro, numero, complemento, bairro, cidade, uf, cep FROM tb_empresa WHERE id=?",
            (st.session_state.empresa_id,)
        ).fetchone()
        socios_rows = conn.execute(
            """SELECT nome, nacionalidade, cpf, rg, orgao_emissor, estado_civil,
            data_nasc, profissao, logradouro, numero, complemento, bairro, cidade, uf, cep
            FROM tb_socio WHERE id_empresa=?""",
            (st.session_state.empresa_id,)
        ).fetchall()

    empresa = {
        'cnpj': emp_row[0], 'nome_empresarial': emp_row[1], 'nome_fantasia': emp_row[2],
        'nire': emp_row[3], 'logradouro': emp_row[4], 'numero': emp_row[5],
        'complemento': emp_row[6], 'bairro': emp_row[7], 'cidade': emp_row[8],
        'uf': emp_row[9], 'cep': emp_row[10]
    }
    socios = [
        {
            'nome': r[0], 'nacionalidade': r[1], 'cpf': r[2], 'rg': r[3],
            'orgao_emissor': r[4], 'estado_civil': r[5], 'data_nasc': r[6],
            'profissao': r[7], 'logradouro': r[8], 'numero': r[9],
            'complemento': r[10], 'bairro': r[11], 'cidade': r[12], 'uf': r[13], 'cep': r[14]
        }
        for r in socios_rows
    ]

    tipo = st.session_state.evento_tipo
    dados = st.session_state.evento_dados

    nomes_tipo = {
        'mudanca_endereco': 'Mudança de Endereço',
        'mudanca_nome': 'Mudança de Nome Empresarial',
        'mudanca_nome_fantasia': 'Mudança de Nome Fantasia'
    }

    # ── Validação: dados obrigatórios para gerar o documento ──
    bloqueios = []

    # Empresa: campos obrigatórios
    campos_emp = {'CNPJ': empresa['cnpj'], 'Logradouro da sede': empresa['logradouro'],
                'Número da sede': empresa['numero'], 'Bairro da sede': empresa['bairro'],
                'Cidade da sede': empresa['cidade'], 'UF da sede': empresa['uf'],
                'CEP da sede': empresa['cep']}
    faltam_emp = [k for k, v in campos_emp.items() if not v]
    if faltam_emp:
        bloqueios.append(f"**Empresa:** {', '.join(faltam_emp)}")

    # Sócios: ao menos um, com campos obrigatórios
    if not socios:
        bloqueios.append("**Sócios:** nenhum sócio cadastrado")
    else:
        for i, s in enumerate(socios, 1):
            campos_s = {'Nome': s['nome'], 'CPF': s['cpf'], 'Nacionalidade': s['nacionalidade'],
                        'Logradouro': s['logradouro'], 'Número': s['numero'],
                        'Bairro': s['bairro'], 'Cidade': s['cidade'], 'UF': s['uf'], 'CEP': s['cep']}
            faltam_s = [k for k, v in campos_s.items() if not v]
            if faltam_s:
                bloqueios.append(f"**Sócio {i} ({s['nome'] or 'sem nome'}):** {', '.join(faltam_s)}")

    if bloqueios:
        st.error("Não é possível gerar o documento. Complete os dados obrigatórios antes de continuar:")
        for b in bloqueios:
            st.markdown(f"- {b}")
        st.markdown("<br>", unsafe_allow_html=True)
        c1, c2 = st.columns(2)
        if c1.button("✏️ EDITAR EMPRESA", type="primary", use_container_width=True):
            nav('editar_empresa')
        if c2.button("⬅️ VOLTAR AOS EVENTOS", type="secondary", use_container_width=True):
            nav('escolher_evento')
        return

    # ── Resumo do documento ──
    st.markdown("### 📋 Resumo do documento")
    with st.expander("Clique para ver o resumo completo antes de baixar", expanded=True):
        st.markdown(f"**Tipo de alteração:** {nomes_tipo.get(tipo, tipo)}")
        st.markdown(f"**Empresa:** {empresa['nome_empresarial']} — CNPJ: {empresa['cnpj']}")
        end_emp = f"{empresa['logradouro']}, nº {empresa['numero']}"
        if empresa['complemento']: end_emp += f", {empresa['complemento']}"
        end_emp += f", {empresa['bairro']}, {empresa['cidade']}/{empresa['uf']} — CEP {empresa['cep']}"
        st.markdown(f"**Sede atual:** {end_emp}")

        st.markdown("**Sócio(s):**")
        for s in socios:
            end_s = f"{s['logradouro']}, nº {s['numero']}"
            if s['complemento']: end_s += f", {s['complemento']}"
            end_s += f", {s['bairro']}, {s['cidade']}/{s['uf']}"
            st.markdown(
                f"- **{s['nome']}** — CPF: {s['cpf']} | {s['nacionalidade']} | "
                f"{s['estado_civil'] or '—'} | Nasc.: {s['data_nasc'] or '—'} | "
                f"Profissão: {s['profissao'] or '—'}\n  Endereço: {end_s}"
            )

        st.markdown("**Alteração proposta:**")
        if tipo == 'mudanca_endereco':
            nd = dados
            novo = f"{nd.get('logradouro','')}, nº {nd.get('numero','')}"
            if nd.get('complemento'): novo += f", {nd['complemento']}"
            novo += f", {nd.get('bairro','')}, {nd.get('cidade','')}/{nd.get('uf','')} — CEP {nd.get('cep','')}"
            st.markdown(f"Nova sede: **{novo}**")
        elif tipo == 'mudanca_nome':
            st.markdown(f"Novo nome empresarial: **{dados.get('nome_empresarial','')}**")
        elif tipo == 'mudanca_nome_fantasia':
            st.markdown(f"Novo nome fantasia: **{dados.get('nome_fantasia','')}**")

    st.markdown("<br>", unsafe_allow_html=True)

    doc_bytes = gerar_doc_alteracao(empresa, socios, tipo, dados)
    nome_arquivo = f"Alteracao_{empresa['nome_empresarial'].replace(' ', '_')}.docx"

    st.download_button(
        label="⬇️ BAIXAR DOCUMENTO .DOCX",
        data=doc_bytes,
        file_name=nome_arquivo,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        type="primary",
        use_container_width=True
    )

    # Registrar no histórico apenas uma vez por visita à tela
    chave_reg = f"_doc_registrado_{tipo}_{st.session_state.empresa_id}"
    if st.session_state.get('user_id') and not st.session_state.get(chave_reg):
        with get_db() as conn:
            conn.execute(
                "INSERT INTO tb_documento (id_empresa, id_usuario, tipo_documento) VALUES (?,?,?)",
                (st.session_state.empresa_id, st.session_state.user_id, nomes_tipo.get(tipo, tipo))
            )
            conn.commit()
        st.session_state[chave_reg] = True

    st.markdown("<br>", unsafe_allow_html=True)
    c1, c2 = st.columns(2)
    if c1.button("⬅️ OUTRO EVENTO", type="secondary", use_container_width=True):
        chave = f"_doc_registrado_{st.session_state.evento_tipo}_{st.session_state.empresa_id}"
        st.session_state.pop(chave, None)
        nav('escolher_evento')
    if c2.button("🏠 MENU PRINCIPAL", type="secondary", use_container_width=True):
        chave = f"_doc_registrado_{st.session_state.evento_tipo}_{st.session_state.empresa_id}"
        st.session_state.pop(chave, None)
        nav('menu')

def tela_historico():
    cabecalho()
    st.subheader("📄 Histórico Geral de Documentos")
    with get_db() as conn:
        docs = conn.execute("""
            SELECT d.data_geracao, d.tipo_documento, e.nome_empresarial, u.nome, d.status
            FROM tb_documento d
            JOIN tb_empresa e ON d.id_empresa = e.id
            JOIN tb_usuario u ON d.id_usuario = u.id
            ORDER BY d.data_geracao DESC
        """).fetchall()
    if not docs:
        st.info("Nenhum documento gerado ainda.")
    else:
        for doc in docs:
            st.markdown(f"**{doc[2]}** — {doc[1]} | Gerado por: {doc[3]} em {doc[0]} | Status: {doc[4]}")
            st.divider()
    if st.button("⬅️ VOLTAR AO MENU", type="secondary"):
        nav('menu')

def tela_historico_empresa():
    cabecalho()
    st.subheader("📄 Histórico da Empresa")
    with get_db() as conn:
        emp = conn.execute("SELECT nome_empresarial FROM tb_empresa WHERE id=?", (st.session_state.empresa_id,)).fetchone()
        docs = conn.execute("""
            SELECT d.data_geracao, d.tipo_documento, u.nome, d.status
            FROM tb_documento d
            JOIN tb_usuario u ON d.id_usuario = u.id
            WHERE d.id_empresa=?
            ORDER BY d.data_geracao DESC
        """, (st.session_state.empresa_id,)).fetchall()
    st.markdown(f"**Empresa:** {emp[0]}")
    if not docs:
        st.info("Nenhum documento gerado para esta empresa.")
    else:
        for doc in docs:
            st.markdown(f"{doc[1]} | Gerado por: {doc[2]} em {doc[0]} | Status: {doc[3]}")
            st.divider()
    if st.button("⬅️ VOLTAR", type="secondary"):
        nav('lista_socios')

# ==========================================
# ROTEADOR
# ==========================================
paginas = {
    'login': tela_login,
    'registrar': tela_registrar,
    'menu': tela_menu,
    'cad_empresa': tela_cad_empresa,
    'editar_empresa': tela_editar_empresa,
    'buscar_empresa': tela_buscar_empresa,
    'lista_socios': tela_socios,
    'cad_socio': tela_cad_socio,
    'editar_socio': tela_editar_socio,
    'escolher_evento': tela_escolher_evento,
    'evento_endereco': tela_evento_endereco,
    'evento_nome': tela_evento_nome,
    'evento_nome_fantasia': tela_evento_nome_fantasia,
    'gerar_documento': tela_gerar_documento,
    'historico': tela_historico,
    'historico_empresa': tela_historico_empresa,
}

paginas.get(st.session_state.pagina, tela_login)()

# Hack de cores para botões ✏️ e 🗑️
def aplicar_hack_cores():
    components.html("""
    <script>
    const botoes = window.parent.document.querySelectorAll('button');
    botoes.forEach(b => {
        const t = b.innerText;
        if (t.includes('✏️')) {
            b.style.setProperty('background-color', '#ffc107', 'important');
            b.style.setProperty('color', '#212529', 'important');
            b.style.setProperty('border-color', '#ffc107', 'important');
        } else if (t.includes('🗑️')) {
            b.style.setProperty('background-color', '#dc3545', 'important');
            b.style.setProperty('border-color', '#dc3545', 'important');
            b.style.setProperty('color', 'white', 'important');
        }
    });
    </script>
    """, height=0, width=0)

aplicar_hack_cores()